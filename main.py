#from BBDD import DatabaseManager
from fastapi import FastAPI
from fastapi.responses import JSONResponse
import os
from fastapi.middleware.cors import CORSMiddleware
from langchain_core.prompts import PromptTemplate
from langchain_openai import OpenAI
import shutil
import subprocess
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import List
from langchain_community.document_loaders import PyPDFLoader
from docx import Document  # Para generar el archivo .docx
# Función para obtener hrefs de una página

import win32com.client
import time

def check_inbox():
    outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
    inbox = outlook.Folders.Item(1).Folders.Item("Inbox")
    messages = inbox.Items
    messages.Sort("[ReceivedTime]", True)
    latest_message = messages.GetLast()
    return latest_message

def process_email(message):
    subject = message.Subject
    body = message.Body
    sender = message.SenderName
    print(f"Subject: {subject}")
    print(f"Sender: {sender}")
    print(f"Body: {body}")

while True:
    latest_email = check_inbox()
    process_email(latest_email)
    time.sleep(60)  # Esperar 1 minuto antes de volver a comprobar

from dotenv import load_dotenv
import os

load_dotenv()  # Carga las variables de entorno desde el archivo .env

openai_api_key = os.getenv("OPENAI_API_KEY")    


app = FastAPI() 

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8080"],  # Cambia esto según sea necesario
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def leer_archivos_vue(directorio):
    archivos_vue = [archivo for archivo in os.listdir(directorio) if archivo.endswith('.vue')]
    contenido_total = ""

    for archivo in archivos_vue:
        ruta_archivo = os.path.join(directorio, archivo)
        with open(ruta_archivo, 'r', encoding='utf-8') as file:
            contenido_total += f"// Contenido de {archivo}:\n"
            contenido_total += file.read() + "\n"

    return contenido_total


def leer_archivos_c(directorio):
    contenido_total=""
    with open(directorio, 'r', encoding='utf-8') as file:
        contenido_total += "// Contenido de main.c:\n"
        contenido_total += file.read() + "\n"

    return contenido_total

pathBBDD="mongodb+srv://fpratsquilez:GTdbIFamKljeyUvs@worldmanagergame.k3rozlq.mongodb.net/?retryWrites=true&w=majority"



from pymongo import MongoClient


class DatabaseManager:
    def __init__(self, db_name):
        self.db_client=MongoClient(db_name).test

    def obtener_BOM(self, BOMname):
        BOM = self.db_client.test.BOMs.find_one({"name": BOMname})
        return BOM
    
    def obtener_BOMs(self):
        BOMS = self.db_client.proyectoIA.find()
        return [doc["name"] for doc in BOMS]
    

db_manager = DatabaseManager(pathBBDD)

# Definir la estructura de cada archivo en docAdicional
class ArchivosHDD(BaseModel):
    name: str
    ruta: str

# Definir la estructura de los datos recibidos en la solicitud
class DataModelGenerarHDD(BaseModel):
    descripcion: str
    BOM: str
    docAdicional: List[ArchivosHDD]

@app.post("/generarHDD")
async def generar_hdd(data: DataModelGenerarHDD):
 # 1. Leer de BBDD la BOM de
    BOM = db_manager.obtener_BOM(data.BOM)
    if not BOM:
        return JSONResponse(content={"error": "BOM no encontrada"}, status_code=404)
    
    # 2. Extraer la ruta de datasheet de cada componente de la BOM
    datasheet_paths = [component['pathdatasheet'] for component in BOM['components']]

    # 3. Leer los pdf de datasheet de cada componente y extraer información relevante
    info_relevante = []
    for path in datasheet_paths:
        loader = PyPDFLoader(path)
        pages = loader.load()
        question = pages
        template = """Question: Extraer información relevante como pinout, descripción del componente, consumo, etc.
        Datasheet: {question}
        Answer: """
        prompt = PromptTemplate.from_template(template)
        llm = OpenAI(openai_api_key=OPENAI_API_KEY, max_tokens=-1, model_name="gpt-3.5-turbo-16k")
        llm_chain = prompt | llm
        info_relevante.append(llm_chain.invoke(question))

    # 4. Leer los pdf de docAdicional
    info_adicional = []
    for doc in data.docAdicional:
        loader = PyPDFLoader(doc.ruta)
        pages = loader.load()
        info_adicional.append(pages)

    # 5. Generar un prompt con toda la información extraída
    combined_info = f"Descripción: {data.descripcion}\n\n"
    combined_info += "Información de Datasheets:\n" + "\n".join(info_relevante) + "\n\n"
    combined_info += "Información de Documentos Adicionales:\n" + "\n".join(info_adicional)
    
    final_template = """Genera un documento de descripción de diseño Hardware con la siguiente información:
    {combined_info}
    El documento debe seguir la estructura estándar de un diseño de hardware.
    """
    final_prompt = PromptTemplate.from_template(final_template)
    final_llm_chain = final_prompt | llm

    # 6. Enviar el prompt a un LLM para generar el documento de descripción de diseño Hardware
    design_description = final_llm_chain.invoke(combined_info)

    # 7. Generar un documento .docx del hardware design description
    doc = Document()
    doc.add_heading('Descripción del Diseño de Hardware', 0)
    doc.add_paragraph(design_description)

    file_path = "C:\Users\xisco\Documents\ProjectManagerIA\HDD IA"  
    doc.save(file_path)

    # 8. Enviar el enlace al front end como respuesta de post
    response = {
        "descripcion": data.descripcion,
        "BOM": data.BOM,
        "docAdicional": data.docAdicional,
        "link_documento": file_path
    }

    return JSONResponse(content=response)

@app.get("/loadBOMs/")
async def loadBOMs():
    BOMS = db_manager.obtener_BOMs()
    return {"mensaje": BOMS}

@app.get("/checkUser/")
async def checkUser():

    return {"mensaje":"No tienes suficiente dinero para mejorar la seguridad"}

@app.get("/showTestBench/")
async def showTestBench():

    return {"mensaje":"No tienes suficiente dinero para mejorar la seguridad"}

@app.get("/startSequence/")
async def startSequence():

    return {"mensaje":""}

@app.get("/modificar")
async def modificar(texto: str):

        # Ejemplo de uso
    directorio = 'C:\\Users\\xisco\\Documents\\ProjectManagerIA\\ProjectManagerWebAI\\FrontEnd\\vue-d3-examples-master\\src\\components\\'
    contenido_components = leer_archivos_vue(directorio)
            # Ejemplo de uso
    directorio = 'C:\\Users\\xisco\\Documents\\ProjectManagerIA\\ProjectManagerWebAI\\FrontEnd\\vue-d3-examples-master\\src\\views\\'
    contenido_view = leer_archivos_vue(directorio)
    contenido_app=""
    ruta_archivo = 'C:\\Users\\xisco\\Documents\\ProjectManagerIA\\ProjectManagerWebAI\\FrontEnd\\vue-d3-examples-master\\src\\App.vue'
    with open(ruta_archivo, 'r', encoding='utf-8') as file:
        contenido_app += f"// Contenido de App.vue:\n"
        contenido_app += file.read() + "\n"
    
    router=""
    ruta_archivo = 'C:\\Users\\xisco\\Documents\\ProjectManagerIA\\ProjectManagerWebAI\\FrontEnd\\vue-d3-examples-master\\src\\router.js'
    with open(ruta_archivo, 'r', encoding='utf-8') as file:
        router += f"// Contenido de router.js:\n"
        router += file.read() + "\n"

    main=""
    ruta_archivo = 'C:\\Users\\xisco\\Documents\\ProjectManagerIA\\ProjectManagerWebAI\\FrontEnd\\vue-d3-examples-master\\src\\main.js'
    with open(ruta_archivo, 'r', encoding='utf-8') as file:
        main += f"// Contenido de main.js:\n"
        main += file.read() + "\n"

    # Rutas de los directorios
    origen = "C:\\Users\\xisco\\Documents\\ProjectManagerIA\\ProjectManagerWebAI\\Pruebas\\HelloWorld.vue"
    destinopath ="C:\\Users\\xisco\\Documents\\ProjectManagerIA\\ProjectManagerWebAI\\FrontEnd\\vue-d3-examples-master\\src\\components\\HelloWorld.vue"
    destino = "C:\\Users\\xisco\\Documents\\ProjectManagerIA\\ProjectManagerWebAI\\FrontEnd\\vue-d3-examples-master\\src\\components"
    
    # # Mover el archivo
    # os.remove(destinopath)
    # shutil.move(origen, destino)
    print(texto)
    question=texto
    template = """Question: Tenemos un programa en vue.js version 3 que esta formado por una carpeta llamada src que contine la carpeta assets, components, plugins y view. En la carpeta components tenemos los archivos: """+contenido_components+"""  En la carpeta view tenemos los archivos: """+contenido_view+"""  En la misma raiz de la carpeta src tenemos el archivo App.vue con el siguiente contenido: """+contenido_app+"""  En la misma raiz de la carpeta src tenemos el archivo router.js con el siguiente contenido: """+router+"""  En la misma raiz de la carpeta src tenemos el archivo main.js con el siguiente contenido: """+main+""". Ahora te voy a dejar la modificacion que quiero que realices sobre el documento que consideres a raiz de la modificacion. Es posible que se indique añadir nuevo codigo. La cuestion es la siguiente: {question}. Es muy importante que me devuelvas el nombre del archivo.vue o .js que vayas a modificar o añadir seguido del codigo. Evita incluir informacion extra.
                Answer:  """ 
    
    # Abre (o crea) el archivo en modo de escritura ('w')
    with open("archivo.txt", "w") as archivo:
        # Escribe el contenido de la variable en el archivo
        archivo.write(template)

    print(template) 
    prompt = PromptTemplate.from_template(template)
    llm = OpenAI(openai_api_key=OPENAI_API_KEY, max_tokens=-1)
    llm_chain = prompt | llm    
    data=llm_chain.invoke(question)

    # Verifica el formato de data
    print(data)

    # Aquí puedes manejar el texto recibido como necesites
    return JSONResponse(content={"message": f"Texto recibido: {data}"})

@app.get("/modificarFW")
async def modificarFW(texto: str):

        # Ejemplo de uso
    directorio = 'C:\\Users\\xisco\\Documents\\ProjectManagerIA\\STM32F429I-Discovery_FW_V1.0.3\\Projects\\Peripheral_Examples\\ADC_DMA\\main.c'
    contenido_components = leer_archivos_c(directorio)
           

    # Rutas de los directorios
    origen = "C:\\Users\\xisco\\Documents\\ProjectManagerIA\\ProjectManagerWebAI\\Pruebas\\main.c"
    destinopath ="C:\\Users\\xisco\\Documents\\ProjectManagerIA\\STM32F429I-Discovery_FW_V1.0.3\\Projects\\Peripheral_Examples\\Touch_Panel\\main.c"
    destino = "C:\\Users\\xisco\\Documents\\ProjectManagerIA\\STM32F429I-Discovery_FW_V1.0.3\\Projects\\Peripheral_Examples\\Touch_Panel\\"
    
    # Mover el archivo
    os.remove(destinopath)
    shutil.move(origen, destino)
    print(texto)

    # Ruta completa del archivo a ejecutar
    ruta_completa = "C:\\Users\\xisco\\Documents\\ProjectManagerIA\\STM32F429I-Discovery_FW_V1.0.3\\Projects\\Peripheral_Examples\\Touch_Panel\\TrueSTUDIO\\Touch_Panel\\Debug\\main.py"

        # Ejecutar el archivo usando subprocess.run()
    # result = subprocess.run(["python", ruta_completa])  
    # print("Salida del script:")
    # print(result.stdout)

    with open(ruta_completa) as f:
        code = f.read()
    exec(code)
    # question=texto
    # template = """Question: Tenemos un programa desarrollado en .c para el microcontrolador stm32f4xx. El main.c es el siguiente: """+contenido_components+""" Ahora te voy a dejar la modificacion que quiero que realices sobre el main.c. Es posible que se indique añadir nuevo codigo o modifiqus el codigo. La cuestion es la siguiente: {question}. Es muy importante que me devuelvas el solo el codigo. Evita incluir informacion extra.
    #             Answer:  """ 
    
    # # Abre (o crea) el archivo en modo de escritura ('w')
    # with open("archivo2.txt", "w") as archivo:
    #     # Escribe el contenido de la variable en el archivo
    #     archivo.write(template)

    # print(template) 
    # prompt = PromptTemplate.from_template(template)
    # llm = OpenAI(openai_api_key=OPENAI_API_KEY, max_tokens=-1)
    # llm_chain = prompt | llm    
    # data=llm_chain.invoke(question)

    data="ok"
    # Verifica el formato de data
    print(data)

    # Aquí puedes manejar el texto recibido como necesites
    return JSONResponse(content={"message": f"Texto recibido: {data}"})

